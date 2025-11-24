# converters/document_converter.py - Fixed LibreOffice detection
import os
import subprocess
import shlex
import json
import csv
import pandas as pd
from typing import Optional
import logging
from io import StringIO, BytesIO
import xlsxwriter
import openpyxl
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

logger = logging.getLogger(__name__)

# Check LibreOffice installation - FIXED VERSION
def check_libreoffice():
    """Check if LibreOffice is installed and accessible"""
    # Possible LibreOffice paths
    possible_paths = [
        'libreoffice',  # System PATH
        'soffice',      # Alternative command
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe'
    ]
    
    for path in possible_paths:
        try:
            print(f"Checking LibreOffice at: {path}")
            
            # Use different approach for Windows
            if os.name == 'nt':
                result = subprocess.run(
                    [path, '--version'], 
                    capture_output=True, 
                    text=True, 
                    timeout=15,
                    creationflags=subprocess.CREATE_NO_WINDOW,
                    shell=False
                )
            else:
                result = subprocess.run(
                    [path, '--version'], 
                    capture_output=True, 
                    text=True, 
                    timeout=15
                )
            
            if result.returncode == 0:
                version_line = result.stdout.split('\n')[0] if result.stdout else "Unknown"
                print(f"[OK] LibreOffice found: {version_line}")
                logger.info(f"LibreOffice available: {version_line}")
                
                # Store the working path for later use
                global LIBREOFFICE_PATH
                LIBREOFFICE_PATH = path
                return True
            else:
                print(f"[FAILED] LibreOffice failed with return code: {result.returncode}")
                if result.stderr:
                    print(f"Error: {result.stderr}")
                
        except subprocess.TimeoutExpired:
            print(f"[FAILED] LibreOffice check timed out for: {path}")
        except FileNotFoundError:
            print(f"[FAILED] LibreOffice not found at: {path}")
        except Exception as e:
            print(f"[FAILED] LibreOffice check error for {path}: {str(e)}")

    print("[FAILED] LibreOffice not found in any standard location")
    return False

# Global variable to store working LibreOffice path
LIBREOFFICE_PATH = None
LIBREOFFICE_AVAILABLE = check_libreoffice()

async def excel_to_pdf(input_path: str, output_path: str) -> bool:
    """Convert Excel files to PDF using LibreOffice"""
    
    print(f"Starting Excel to PDF conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")
    
    if not LIBREOFFICE_AVAILABLE or not LIBREOFFICE_PATH:
        print("[FAILED] LibreOffice not available - cannot convert Excel to PDF")
        return False
    
    if not os.path.exists(input_path):
        print(f"[FAILED] Input file does not exist: {input_path}")
        return False
    
    try:
        # Use LibreOffice headless mode for conversion
        output_dir = os.path.dirname(output_path)
        
        cmd = [
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            input_path
        ]
        
        print(f"LibreOffice command: {' '.join(cmd)}")
        
        if os.name == 'nt':
            process = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300,  # 5 minutes
                creationflags=subprocess.CREATE_NO_WINDOW,
                shell=False
            )
        else:
            process = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300
            )
        
        if process.returncode == 0:
            # LibreOffice creates PDF with same name as input
            input_name = os.path.splitext(os.path.basename(input_path))[0]
            generated_pdf = os.path.join(output_dir, f"{input_name}.pdf")
            
            if os.path.exists(generated_pdf):
                # Rename to desired output name if different
                if generated_pdf != output_path:
                    os.rename(generated_pdf, output_path)
                
                output_size = os.path.getsize(output_path)
                print(f"✅ Excel to PDF conversion successful!")
                print(f"Output file size: {output_size} bytes")
                return True
            else:
                print("[FAILED] PDF output file not found")
                print(f"Expected: {generated_pdf}")
                return False
        else:
            print(f"[FAILED] LibreOffice failed: {process.stderr}")
            print(f"stdout: {process.stdout}")
            return False
            
    except subprocess.TimeoutExpired:
        print("[FAILED] Excel to PDF conversion timed out")
        return False
    except Exception as e:
        print(f"[FAILED] Excel to PDF conversion error: {str(e)}")
        return False

async def word_to_pdf(input_path: str, output_path: str) -> bool:
    """Convert Word documents (DOCX, DOC) to PDF using pypandoc (best quality)"""

    print(f"Starting Word to PDF conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")

    if not os.path.exists(input_path):
        print(f"[FAILED] Input file does not exist: {input_path}")
        return False

    try:
        # Method 1: Try pypandoc (best quality, universal document converter)
        try:
            import pypandoc
            print("Using pypandoc (best quality converter)...")

            # Convert using pypandoc - try multiple PDF engines
            pdf_engines = ['wkhtmltopdf', 'weasyprint', 'pdflatex', 'xelatex']

            for engine in pdf_engines:
                try:
                    print(f"Trying pypandoc with {engine}...")
                    pypandoc.convert_file(
                        input_path,
                        'pdf',
                        outputfile=output_path,
                        extra_args=[f'--pdf-engine={engine}']
                    )

                    if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                        output_size = os.path.getsize(output_path)
                        print(f"[SUCCESS] Word to PDF conversion successful (pypandoc with {engine})!")
                        print(f"Output file size: {output_size} bytes")
                        return True
                except Exception as engine_error:
                    print(f"pypandoc with {engine} failed: {str(engine_error)}")
                    continue

        except Exception as pypandoc_error:
            print(f"pypandoc failed: {str(pypandoc_error)}")
            print("Falling back to docx2pdf...")

        # Method 2: Try using docx2pdf (Windows COM - uses Microsoft Word if available)
        if os.name == 'nt':
            try:
                from docx2pdf import convert as docx2pdf_convert
                print("Using docx2pdf with Microsoft Word COM automation...")
                docx2pdf_convert(input_path, output_path)

                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    output_size = os.path.getsize(output_path)
                    print(f"[SUCCESS] Word to PDF conversion successful (docx2pdf)!")
                    print(f"Output file size: {output_size} bytes")
                    return True
            except Exception as docx2pdf_error:
                print(f"docx2pdf failed: {str(docx2pdf_error)}")
                print("Falling back to python-docx + reportlab...")

        # Method 3: Use python-docx + reportlab (cross-platform fallback)
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

        print("Using python-docx + reportlab...")

        doc = Document(input_path)
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
        heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, spaceAfter=10)
        normal_style = ParagraphStyle('Normal2', parent=styles['Normal'], fontSize=11, spaceAfter=8, leading=14)

        # Convert paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 0.2*inch))
                continue

            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            if para.style.name.startswith('Heading 1') or para.style.name == 'Title':
                story.append(Paragraph(text, title_style))
            elif para.style.name.startswith('Heading'):
                story.append(Paragraph(text, heading_style))
            else:
                story.append(Paragraph(text, normal_style))

        # Convert tables
        for table in doc.tables:
            story.append(Spacer(1, 0.2*inch))
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                if row_text.strip():
                    story.append(Paragraph(row_text, normal_style))
            story.append(Spacer(1, 0.2*inch))

        if not story:
            story.append(Paragraph("Document converted from Word", normal_style))

        pdf.build(story)

        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            output_size = os.path.getsize(output_path)
            print(f"[SUCCESS] Word to PDF conversion successful (reportlab)!")
            print(f"Output file size: {output_size} bytes")
            return True
        else:
            print("[FAILED] PDF not created")
            return False

    except Exception as e:
        print(f"[FAILED] Word to PDF error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

async def powerpoint_to_pdf(input_path: str, output_path: str) -> bool:
    """Convert PowerPoint files to PDF using LibreOffice"""
    
    print(f"Starting PowerPoint to PDF conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")
    
    if not LIBREOFFICE_AVAILABLE or not LIBREOFFICE_PATH:
        print("[FAILED] LibreOffice not available - cannot convert PowerPoint to PDF")
        return False
    
    try:
        output_dir = os.path.dirname(output_path)
        
        cmd = [
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            input_path
        ]
        
        if os.name == 'nt':
            process = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300,
                creationflags=subprocess.CREATE_NO_WINDOW,
                shell=False
            )
        else:
            process = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300
            )
        
        if process.returncode == 0:
            input_name = os.path.splitext(os.path.basename(input_path))[0]
            generated_pdf = os.path.join(output_dir, f"{input_name}.pdf")
            
            if os.path.exists(generated_pdf):
                if generated_pdf != output_path:
                    os.rename(generated_pdf, output_path)
                
                print(f"✅ PowerPoint to PDF conversion successful!")
                return True
            else:
                print("[FAILED] PDF output file not found")
                return False
        else:
            print(f"[FAILED] LibreOffice failed: {process.stderr}")
            return False
            
    except Exception as e:
        print(f"[FAILED] PowerPoint to PDF conversion error: {str(e)}")
        return False

async def text_to_pdf(input_path: str, output_path: str) -> bool:
    """Convert text files to PDF using ReportLab"""
    
    print(f"Starting Text to PDF conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")
    
    try:
        # Read text file
        with open(input_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Create a custom style for better text formatting
        custom_style = ParagraphStyle(
            'CustomStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            spaceAfter=6,
        )
        
        # Split text into paragraphs and create PDF content
        story = []
        paragraphs = text_content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                # Replace single newlines with <br/> for line breaks
                formatted_para = para.replace('\n', '<br/>')
                story.append(Paragraph(formatted_para, custom_style))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"✅ Text to PDF conversion successful!")
            return True
        else:
            print("[FAILED] PDF creation failed")
            return False
            
    except Exception as e:
        print(f"[FAILED] Text to PDF conversion error: {str(e)}")
        return False

async def html_to_pdf(input_path: str, output_path: str) -> bool:
    """Convert HTML files to PDF using wkhtmltopdf or LibreOffice fallback"""
    
    print(f"Starting HTML to PDF conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")
    
    try:
        # Try wkhtmltopdf first
        try:
            import pdfkit
            
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None
            }
            
            # Read HTML content
            with open(input_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Convert to PDF
            pdfkit.from_string(html_content, output_path, options=options)
            
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"✅ HTML to PDF conversion successful (wkhtmltopdf)!")
                return True
                
        except Exception as e:
            print(f"wkhtmltopdf failed: {str(e)}")
            
        # Fallback to LibreOffice if available
        if LIBREOFFICE_AVAILABLE and LIBREOFFICE_PATH:
            print("Trying LibreOffice fallback...")
            
            output_dir = os.path.dirname(output_path)
            cmd = [
                LIBREOFFICE_PATH,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                input_path
            ]
            
            if os.name == 'nt':
                process = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=300,
                    creationflags=subprocess.CREATE_NO_WINDOW,
                    shell=False
                )
            else:
                process = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if process.returncode == 0:
                input_name = os.path.splitext(os.path.basename(input_path))[0]
                generated_pdf = os.path.join(output_dir, f"{input_name}.pdf")
                
                if os.path.exists(generated_pdf):
                    if generated_pdf != output_path:
                        os.rename(generated_pdf, output_path)
                    print(f"✅ HTML to PDF conversion successful (LibreOffice)!")
                    return True
        
        print("[FAILED] HTML to PDF conversion failed - no converter available")
        return False
            
    except Exception as e:
        print(f"[FAILED] HTML to PDF conversion error: {str(e)}")
        return False

async def csv_to_excel(input_path: str, output_path: str) -> bool:
    """Convert CSV files to Excel using pandas and openpyxl"""
    
    print(f"Starting CSV to Excel conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")
    
    try:
        # Read CSV file
        df = pd.read_csv(input_path, encoding='utf-8')
        
        # Write to Excel
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Sheet1', index=False)
            
            # Auto-adjust column widths
            worksheet = writer.sheets['Sheet1']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"✅ CSV to Excel conversion successful!")
            return True
        else:
            print("[FAILED] Excel creation failed")
            return False
            
    except Exception as e:
        print(f"[FAILED] CSV to Excel conversion error: {str(e)}")
        return False

async def json_to_csv(input_path: str, output_path: str) -> bool:
    """Convert JSON files to CSV using pandas"""

    print(f"Starting JSON to CSV conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")

    try:
        # Read JSON file
        with open(input_path, 'r', encoding='utf-8') as file:
            json_data = json.load(file)

        # Handle different JSON structures
        if isinstance(json_data, list):
            # Array of objects
            df = pd.json_normalize(json_data)
        elif isinstance(json_data, dict):
            # Single object or nested structure
            df = pd.json_normalize([json_data])
        else:
            # Simple value
            df = pd.DataFrame([{'value': json_data}])

        # Write to CSV with UTF-8 BOM for Windows compatibility
        df.to_csv(output_path, index=False, encoding='utf-8-sig')

        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print("[OK] JSON to CSV conversion successful!")
            return True
        else:
            print("[FAILED] CSV creation failed")
            return False

    except Exception as e:
        print(f"[FAILED] JSON to CSV conversion error: {str(e)}")
        return False

async def pdf_to_text(input_path: str, output_path: str) -> bool:
    """Convert PDF files to text using PyPDF2"""

    print(f"Starting PDF to Text conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")

    try:
        import PyPDF2

        # Open the PDF file
        with open(input_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                print("[FAILED] PDF is encrypted and cannot be extracted")
                return False

            # Extract text from all pages
            text_content = []
            total_pages = len(pdf_reader.pages)

            print(f"Extracting text from {total_pages} pages...")

            for page_num in range(total_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()

                if page_text and page_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n")
                    text_content.append(page_text)
                    text_content.append("\n\n")

            # Combine all text
            full_text = "".join(text_content)

            if not full_text.strip():
                print("[FAILED] No text could be extracted from PDF (might be image-based)")
                return False

            # Write to output file
            with open(output_path, 'w', encoding='utf-8') as text_file:
                text_file.write(full_text)

            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"✅ PDF to Text conversion successful!")
                print(f"Extracted {len(full_text)} characters from {total_pages} pages")
                return True
            else:
                print("[FAILED] Text file creation failed")
                return False

    except Exception as e:
        print(f"[FAILED] PDF to Text conversion error: {str(e)}")
        return False

async def pdf_to_word(input_path: str, output_path: str) -> bool:
    """Convert PDF to Word using pdf2docx library"""

    print(f"Starting PDF to Word conversion:")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")

    try:
        from converters.pdf_to_word_converter import pdf_to_word_converter

        # Use the existing converter
        success = await pdf_to_word_converter(input_path, output_path)

        if success and os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"✅ PDF to Word conversion successful!")
            return True
        else:
            print("[FAILED] PDF to Word conversion failed")
            return False

    except Exception as e:
        print(f"[FAILED] PDF to Word conversion error: {str(e)}")
        return False

async def word_to_html(input_path: str, output_path: str) -> bool:
    """Convert Word documents (DOCX) to HTML with EXACT formatting preservation"""

    print(f"Starting Word to HTML conversion (Full Fidelity):")
    print(f"   Input: {input_path}")
    print(f"   Output: {output_path}")

    if not os.path.exists(input_path):
        print(f"[FAILED] Input file does not exist: {input_path}")
        return False

    try:
        # Method 1: Try pypandoc (best quality with full styling)
        try:
            import pypandoc
            print("Using pypandoc for DOCX to HTML conversion...")

            output = pypandoc.convert_file(
                input_path,
                'html',
                outputfile=output_path,
                extra_args=['--standalone', '--embed-resources', '--self-contained']
            )

            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                output_size = os.path.getsize(output_path)
                print(f"[SUCCESS] Word to HTML conversion successful (pypandoc)!")
                print(f"Output file size: {output_size} bytes")
                return True

        except Exception as pypandoc_error:
            print(f"pypandoc failed: {str(pypandoc_error)}")
            print("Falling back to python-docx with full fidelity conversion...")

        # Method 2: Full fidelity conversion using python-docx
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement
        import base64
        import re
        from zipfile import ZipFile

        print("Using python-docx for full fidelity DOCX to HTML conversion...")

        doc = Document(input_path)

        # Extract images from docx
        images_data = {}
        try:
            with ZipFile(input_path, 'r') as zip_file:
                for name in zip_file.namelist():
                    if name.startswith('word/media/'):
                        image_name = name.split('/')[-1]
                        image_data = zip_file.read(name)
                        ext = image_name.split('.')[-1].lower()
                        if ext == 'jpg':
                            ext = 'jpeg'
                        images_data[image_name] = {
                            'data': base64.b64encode(image_data).decode('utf-8'),
                            'type': ext
                        }
        except Exception as e:
            print(f"Warning: Could not extract images: {e}")

        # Build relationship ID to image mapping
        rel_to_image = {}
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in str(rel.target_ref):
                    image_name = str(rel.target_ref).split('/')[-1]
                    if image_name in images_data:
                        rel_to_image[rel_id] = images_data[image_name]
        except:
            pass

        # Start HTML with comprehensive styles matching Word defaults
        html_parts = ['''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        @page {
            size: A4;
            margin: 1in;
        }
        body {
            font-family: 'Calibri', 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.15;
            margin: 0;
            padding: 1in;
            color: #000000;
            background: #fff;
        }
        * {
            box-sizing: border-box;
        }
        p {
            margin: 0;
            padding: 0;
        }
        h1 { font-size: 16pt; font-weight: bold; margin: 12pt 0 3pt 0; }
        h2 { font-size: 13pt; font-weight: bold; margin: 12pt 0 3pt 0; }
        h3 { font-size: 12pt; font-weight: bold; margin: 12pt 0 3pt 0; }
        h4 { font-size: 11pt; font-weight: bold; font-style: italic; margin: 12pt 0 3pt 0; }
        h5 { font-size: 11pt; font-weight: normal; margin: 12pt 0 3pt 0; }
        h6 { font-size: 11pt; font-weight: normal; font-style: italic; margin: 12pt 0 3pt 0; }
        table {
            border-collapse: collapse;
            margin: 0;
            width: auto;
        }
        td, th {
            padding: 5px 7px;
            vertical-align: top;
        }
        ul, ol {
            margin: 0;
            padding-left: 0.5in;
        }
        li {
            margin: 0;
        }
        img {
            max-width: 100%;
        }
        a {
            color: #0563C1;
            text-decoration: underline;
        }
        .page-break {
            page-break-after: always;
        }
    </style>
</head>
<body>
''']

        def get_highlight_color(highlight):
            """Convert Word highlight color to CSS"""
            highlight_colors = {
                'YELLOW': '#FFFF00', 'GREEN': '#00FF00', 'CYAN': '#00FFFF',
                'MAGENTA': '#FF00FF', 'BLUE': '#0000FF', 'RED': '#FF0000',
                'DARK_BLUE': '#000080', 'DARK_CYAN': '#008080', 'DARK_GREEN': '#008000',
                'DARK_MAGENTA': '#800080', 'DARK_RED': '#800000', 'DARK_YELLOW': '#808000',
                'GRAY_25': '#C0C0C0', 'GRAY_50': '#808080', 'BLACK': '#000000',
                'WHITE': '#FFFFFF'
            }
            return highlight_colors.get(str(highlight), '#FFFF00')

        def get_run_html(run):
            """Convert a run to HTML with EXACT styling preservation"""
            text = run.text
            if not text:
                return ''

            # Escape HTML characters
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            text = text.replace('\n', '<br>')
            text = text.replace('\t', '&emsp;')

            # Build comprehensive style string
            styles = []

            # Font family
            if run.font.name:
                styles.append(f"font-family: '{run.font.name}', sans-serif")

            # Font size
            if run.font.size:
                styles.append(f"font-size: {run.font.size.pt}pt")

            # Font color
            if run.font.color and run.font.color.rgb:
                styles.append(f"color: #{run.font.color.rgb}")
            elif run.font.color and run.font.color.theme_color:
                # Handle theme colors
                pass

            # Highlight/background color
            if run.font.highlight_color:
                styles.append(f"background-color: {get_highlight_color(run.font.highlight_color)}")

            # Letter spacing
            try:
                if run.font._element.rPr is not None:
                    spacing = run.font._element.rPr.find(qn('w:spacing'))
                    if spacing is not None and spacing.get(qn('w:val')):
                        val = int(spacing.get(qn('w:val'))) / 20  # twips to pt
                        styles.append(f"letter-spacing: {val}pt")
            except:
                pass

            # Build inline style
            style_str = '; '.join(styles)

            # Apply formatting tags (order matters)
            if run.font.subscript:
                text = f'<sub>{text}</sub>'
            if run.font.superscript:
                text = f'<sup>{text}</sup>'
            if run.font.strike:
                text = f'<s>{text}</s>'
            if run.font.double_strike:
                text = f'<s style="text-decoration-style: double">{text}</s>'
            if run.underline:
                underline_style = 'underline'
                if hasattr(run.font, 'underline') and run.font.underline:
                    u_type = str(run.font.underline)
                    if 'DOUBLE' in u_type:
                        underline_style = 'underline double'
                    elif 'DOTTED' in u_type:
                        underline_style = 'underline dotted'
                    elif 'DASH' in u_type:
                        underline_style = 'underline dashed'
                    elif 'WAVE' in u_type:
                        underline_style = 'underline wavy'
                text = f'<span style="text-decoration: {underline_style}">{text}</span>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.bold:
                text = f'<strong>{text}</strong>'

            # Wrap with style span if needed
            if style_str:
                text = f'<span style="{style_str}">{text}</span>'

            return text

        def get_paragraph_style(para):
            """Get comprehensive paragraph CSS styles"""
            styles = []

            # Alignment
            if para.alignment:
                align_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: 'left',
                    WD_ALIGN_PARAGRAPH.CENTER: 'center',
                    WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                    WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
                    WD_ALIGN_PARAGRAPH.DISTRIBUTE: 'justify',
                }
                if para.alignment in align_map:
                    styles.append(f"text-align: {align_map[para.alignment]}")

            pf = para.paragraph_format

            # Line spacing
            if pf.line_spacing:
                if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                    styles.append(f"line-height: {pf.line_spacing.pt}pt")
                elif pf.line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                    styles.append(f"line-height: {pf.line_spacing.pt}pt")
                elif isinstance(pf.line_spacing, float):
                    styles.append(f"line-height: {pf.line_spacing}")

            # Indentation
            if pf.left_indent:
                styles.append(f"margin-left: {pf.left_indent.pt}pt")
            if pf.right_indent:
                styles.append(f"margin-right: {pf.right_indent.pt}pt")
            if pf.first_line_indent:
                if pf.first_line_indent.pt > 0:
                    styles.append(f"text-indent: {pf.first_line_indent.pt}pt")
                else:
                    # Hanging indent
                    styles.append(f"text-indent: {pf.first_line_indent.pt}pt")
                    styles.append(f"padding-left: {abs(pf.first_line_indent.pt)}pt")

            # Spacing
            if pf.space_before:
                styles.append(f"margin-top: {pf.space_before.pt}pt")
            if pf.space_after:
                styles.append(f"margin-bottom: {pf.space_after.pt}pt")

            # Background/shading
            try:
                shading = para._element.pPr.find(qn('w:shd')) if para._element.pPr else None
                if shading is not None:
                    fill = shading.get(qn('w:fill'))
                    if fill and fill != 'auto':
                        styles.append(f"background-color: #{fill}")
            except:
                pass

            # Borders
            try:
                pBdr = para._element.pPr.find(qn('w:pBdr')) if para._element.pPr else None
                if pBdr is not None:
                    for side in ['top', 'bottom', 'left', 'right']:
                        border = pBdr.find(qn(f'w:{side}'))
                        if border is not None:
                            val = border.get(qn('w:val'))
                            sz = border.get(qn('w:sz'))
                            color = border.get(qn('w:color'))
                            if val and val != 'nil':
                                width = int(sz) / 8 if sz else 1
                                border_color = f'#{color}' if color and color != 'auto' else '#000'
                                styles.append(f"border-{side}: {width}pt solid {border_color}")
            except:
                pass

            return '; '.join(styles)

        def get_paragraph_html(para):
            """Convert a paragraph to HTML with FULL styling preservation"""
            # Get paragraph content with run styling
            content = ''
            for run in para.runs:
                content += get_run_html(run)

            # Check for images in paragraph
            for run in para.runs:
                try:
                    inline_shapes = run._element.findall('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                    for blip in inline_shapes:
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and embed_id in rel_to_image:
                            img_info = rel_to_image[embed_id]
                            content += f'<img src="data:image/{img_info["type"]};base64,{img_info["data"]}" alt="Image">'
                except:
                    pass

            if not content and not para.text.strip():
                return '<p>&nbsp;</p>'

            if not content:
                content = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # Determine tag based on style
            style_name = para.style.name if para.style else 'Normal'

            if style_name == 'Title':
                tag = 'h1'
            elif style_name.startswith('Heading 1'):
                tag = 'h1'
            elif style_name.startswith('Heading 2'):
                tag = 'h2'
            elif style_name.startswith('Heading 3'):
                tag = 'h3'
            elif style_name.startswith('Heading 4'):
                tag = 'h4'
            elif style_name.startswith('Heading 5'):
                tag = 'h5'
            elif style_name.startswith('Heading 6'):
                tag = 'h6'
            elif 'List' in style_name:
                return f'<li style="{get_paragraph_style(para)}">{content}</li>'
            else:
                tag = 'p'

            # Get full paragraph styling
            style_attr = get_paragraph_style(para)
            style_str = f' style="{style_attr}"' if style_attr else ''

            return f'<{tag}{style_str}>{content}</{tag}>'

        def get_table_cell_style(cell):
            """Get table cell CSS styles"""
            styles = []

            # Vertical alignment
            try:
                if cell.vertical_alignment:
                    v_align = {
                        WD_ALIGN_VERTICAL.TOP: 'top',
                        WD_ALIGN_VERTICAL.CENTER: 'middle',
                        WD_ALIGN_VERTICAL.BOTTOM: 'bottom',
                    }
                    if cell.vertical_alignment in v_align:
                        styles.append(f"vertical-align: {v_align[cell.vertical_alignment]}")
            except:
                pass

            # Cell width
            try:
                if cell.width:
                    styles.append(f"width: {cell.width.inches}in")
            except:
                pass

            # Cell shading/background
            try:
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill != 'auto':
                            styles.append(f"background-color: #{fill}")
            except:
                pass

            # Cell borders
            try:
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        for side in ['top', 'bottom', 'left', 'right']:
                            border = tcBorders.find(qn(f'w:{side}'))
                            if border is not None:
                                val = border.get(qn('w:val'))
                                sz = border.get(qn('w:sz'))
                                color = border.get(qn('w:color'))
                                if val and val != 'nil':
                                    width = int(sz) / 8 if sz else 1
                                    border_color = f'#{color}' if color and color != 'auto' else '#000'
                                    styles.append(f"border-{side}: {width}pt solid {border_color}")
            except:
                pass

            return '; '.join(styles)

        def process_table(table):
            """Process table with full formatting"""
            table_styles = []

            # Table width
            try:
                if table._tbl.tblPr is not None:
                    tblW = table._tbl.tblPr.find(qn('w:tblW'))
                    if tblW is not None:
                        w_type = tblW.get(qn('w:type'))
                        w_val = tblW.get(qn('w:w'))
                        if w_type == 'pct':
                            table_styles.append(f"width: {int(w_val)/50}%")
                        elif w_type == 'dxa' and w_val:
                            table_styles.append(f"width: {int(w_val)/20}pt")
            except:
                pass

            style_str = f' style="{"; ".join(table_styles)}"' if table_styles else ''
            html = f'<table{style_str}>'

            for row in table.rows:
                html += '<tr>'
                for cell in row.cells:
                    cell_style = get_table_cell_style(cell)
                    style_attr = f' style="{cell_style}"' if cell_style else ''

                    # Get cell content with formatting
                    cell_content = ''
                    for para in cell.paragraphs:
                        if cell_content:
                            cell_content += '<br>'
                        # Get paragraph content
                        para_content = ''
                        for run in para.runs:
                            para_content += get_run_html(run)
                        if not para_content:
                            para_content = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        cell_content += para_content

                    html += f'<td{style_attr}>{cell_content}</td>'
                html += '</tr>'
            html += '</table>'
            return html

        # Determine if paragraph has numbering (list)
        def is_list_para(para):
            """Check if paragraph is a list item"""
            try:
                numPr = para._element.pPr.find(qn('w:numPr')) if para._element.pPr else None
                return numPr is not None
            except:
                return False

        def get_list_level(para):
            """Get list indentation level"""
            try:
                numPr = para._element.pPr.find(qn('w:numPr')) if para._element.pPr else None
                if numPr is not None:
                    ilvl = numPr.find(qn('w:ilvl'))
                    if ilvl is not None:
                        return int(ilvl.get(qn('w:val')))
            except:
                pass
            return 0

        # Process document body with proper ordering
        # Track element positions for proper ordering
        body_element = doc.element.body

        in_list = False
        list_type = None
        list_level = 0

        for child in body_element:
            if child.tag == qn('w:p'):
                # Find corresponding paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is child:
                        para = p
                        break

                if para is None:
                    continue

                style_name = para.style.name if para.style else 'Normal'

                # Check for page break
                try:
                    if para._element.pPr is not None:
                        pageBreak = para._element.pPr.find(qn('w:pageBreakBefore'))
                        if pageBreak is not None:
                            html_parts.append('<div class="page-break"></div>')
                except:
                    pass

                # Handle lists
                if is_list_para(para) or 'List' in style_name:
                    current_level = get_list_level(para)
                    is_numbered = 'Number' in style_name or any(c.isdigit() for c in style_name)
                    new_list_type = 'ol' if is_numbered else 'ul'

                    if not in_list:
                        html_parts.append(f'<{new_list_type}>')
                        in_list = True
                        list_type = new_list_type
                        list_level = current_level
                    elif new_list_type != list_type:
                        html_parts.append(f'</{list_type}>')
                        html_parts.append(f'<{new_list_type}>')
                        list_type = new_list_type

                    html_parts.append(get_paragraph_html(para))
                else:
                    if in_list:
                        html_parts.append(f'</{list_type}>')
                        in_list = False
                        list_type = None
                    html_parts.append(get_paragraph_html(para))

            elif child.tag == qn('w:tbl'):
                # Close any open list
                if in_list:
                    html_parts.append(f'</{list_type}>')
                    in_list = False
                    list_type = None

                # Find corresponding table
                for table in doc.tables:
                    if table._tbl is child:
                        html_parts.append(process_table(table))
                        break

        # Close any remaining list
        if in_list:
            html_parts.append(f'</{list_type}>')

        html_parts.append('</body>\n</html>')

        with open(output_path, 'w', encoding='utf-8') as html_file:
            html_file.write('\n'.join(html_parts))

        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            output_size = os.path.getsize(output_path)
            print(f"[SUCCESS] Word to HTML conversion successful (python-docx with full styling)!")
            print(f"Output file size: {output_size} bytes")
            return True
        else:
            print("[FAILED] HTML not created")
            return False

    except Exception as e:
        print(f"[FAILED] Word to HTML error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# Supported document formats
SUPPORTED_DOCUMENT_FORMATS = {
    'word_to_pdf': {
        'input': ['docx', 'doc'],
        'output': ['pdf']
    },
    'excel_to_pdf': {
        'input': ['xlsx', 'xls', 'xlsm'],
        'output': ['pdf']
    },
    'powerpoint_to_pdf': {
        'input': ['pptx', 'ppt', 'pptm'],
        'output': ['pdf']
    },
    'text_to_pdf': {
        'input': ['txt', 'text'],
        'output': ['pdf']
    },
    'pdf_to_text': {
        'input': ['pdf'],
        'output': ['txt', 'text']
    },
    'pdf_to_word': {
        'input': ['pdf'],
        'output': ['docx', 'doc']
    },
    'html_to_pdf': {
        'input': ['html', 'htm'],
        'output': ['pdf']
    },
    'csv_to_excel': {
        'input': ['csv'],
        'output': ['xlsx']
    },
    'json_to_csv': {
        'input': ['json'],
        'output': ['csv']
    },
    'word_to_html': {
        'input': ['docx', 'doc'],
        'output': ['html', 'htm']
    }
}