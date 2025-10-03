# converters/docx_to_epub_converter.py - DOCX to EPUB conversion logic
from utils.dependencies import DOCX_AVAILABLE, EPUB_AVAILABLE

async def docx_to_epub_converter(docx_path: str, output_path: str) -> bool:
    """Convert DOCX to EPUB document"""
    if not DOCX_AVAILABLE or not EPUB_AVAILABLE:
        return False

    try:
        from docx import Document
        import ebooklib
        from ebooklib import epub
        from pathlib import Path

        # Read DOCX file
        doc = Document(docx_path)

        # Create EPUB book
        book = epub.EpubBook()

        # Set metadata
        filename = Path(docx_path).stem
        book.set_identifier(f'docx_converted_{filename}')
        book.set_title(filename.replace('_', ' ').title())
        book.set_language('en')
        book.add_author('Unknown')

        # Extract content from DOCX
        chapters = []
        current_chapter = []
        chapter_title = "Chapter 1"
        chapter_count = 1

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if not text:
                continue

            # Check if this is a heading (potential chapter break)
            if paragraph.style.name.startswith('Heading'):
                # Save previous chapter if it has content
                if current_chapter:
                    chapters.append((chapter_title, current_chapter))

                # Start new chapter
                chapter_title = text if text else f"Chapter {chapter_count + 1}"
                current_chapter = []
                chapter_count += 1
            else:
                # Add paragraph to current chapter
                current_chapter.append(paragraph)

        # Add final chapter
        if current_chapter:
            chapters.append((chapter_title, current_chapter))

        # If no chapters were created (no headings), create one chapter with all content
        if not chapters:
            all_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            if all_paragraphs:
                chapters.append(("Full Document", all_paragraphs))

        # Create EPUB chapters
        epub_chapters = []
        for i, (title, paragraphs) in enumerate(chapters):
            # Create HTML content
            html_content = f"""
            <html>
            <head>
                <title>{title}</title>
                <style>
                    body {{ font-family: Georgia, serif; line-height: 1.6; margin: 2em; }}
                    h1 {{ color: #333; border-bottom: 2px solid #333; }}
                    p {{ margin-bottom: 1em; text-align: justify; }}
                    .bold {{ font-weight: bold; }}
                    .italic {{ font-style: italic; }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
            """

            # Process paragraphs
            for para in paragraphs:
                if para.text.strip():
                    # Basic formatting preservation
                    para_html = "<p>"
                    for run in para.runs:
                        text = run.text
                        if run.bold and run.italic:
                            text = f"<strong><em>{text}</em></strong>"
                        elif run.bold:
                            text = f"<strong>{text}</strong>"
                        elif run.italic:
                            text = f"<em>{text}</em>"
                        para_html += text
                    para_html += "</p>"
                    html_content += para_html

            html_content += """
            </body>
            </html>
            """

            # Create EPUB chapter
            chapter = epub.EpubHtml(
                title=title,
                file_name=f'chap_{i:02d}.xhtml',
                lang='en'
            )
            chapter.content = html_content

            book.add_item(chapter)
            epub_chapters.append(chapter)

        # Create table of contents
        book.toc = tuple(epub.Link(f"chap_{i:02d}.xhtml", title, f"chapter{i}")
                        for i, (title, _) in enumerate(chapters))

        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Create spine
        book.spine = ['nav'] + epub_chapters

        # Write EPUB file
        epub.write_epub(output_path, book, {})
        return True

    except Exception as e:
        print(f"DOCX to EPUB conversion error: {e}")
        return False