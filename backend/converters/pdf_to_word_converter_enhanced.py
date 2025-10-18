# converters/pdf_to_word_converter_enhanced.py
"""
Enhanced PDF to Word converter with multiple conversion methods

This module provides advanced PDF to Word conversion with support for:
- Multiple conversion engines (pdf2docx, PyMuPDF, pdfplumber)
- Automatic fallback mechanism
- Color and styling preservation
- Layout and positioning accuracy
- Method selection via API parameter

Author: Backend Team
Date: October 2025
"""

import os
from typing import Optional, Tuple
from utils.dependencies import PDF2DOCX_AVAILABLE

# Import conversion methods
from .pdf_to_word_methods.pdf2docx_method import pdf2docx_convert
from .pdf_to_word_methods.pymupdf_method import pymupdf_to_word
from .pdf_to_word_methods.pdfplumber_method import pdfplumber_to_word_advanced


async def pdf_to_word_converter_enhanced(
    pdf_path: str,
    output_path: str,
    method: Optional[str] = "auto"
) -> Tuple[bool, str]:
    """
    Enhanced PDF to Word converter with multiple methods

    Args:
        pdf_path: Input PDF file path
        output_path: Output DOCX file path
        method: Conversion method to use
            - "auto": Try all methods in order of quality (default)
            - "pdf2docx": Use pdf2docx library (best for general use)
            - "pymupdf": Use PyMuPDF with styling extraction (best for colors)
            - "pdfplumber": Use pdfplumber with positioning (best for layout)

    Returns:
        Tuple[bool, str]: (success, method_used)
            - success: True if conversion succeeded, False otherwise
            - method_used: Name of the method that succeeded, or error message

    Example:
        success, method = await pdf_to_word_converter_enhanced(
            "input.pdf",
            "output.docx",
            method="auto"
        )
        if success:
            print(f"Converted successfully using {method}")
    """

    print(f"[INFO] PDF to Word conversion requested with method: {method}")
    print(f"[INFO] Input: {pdf_path}")
    print(f"[INFO] Output: {output_path}")

    if method == "auto":
        # Try methods in order of quality and reliability
        methods = [
            ("pdf2docx", pdf2docx_convert, PDF2DOCX_AVAILABLE),
            ("pymupdf", pymupdf_to_word, True),  # PyMuPDF is usually available
            ("pdfplumber", pdfplumber_to_word_advanced, True),  # pdfplumber is usually available
        ]

        print("[INFO] Auto mode: Trying multiple conversion methods...")

        for method_name, method_func, available in methods:
            if not available:
                print(f"[SKIP] {method_name} is not available, skipping...")
                continue

            try:
                print(f"[ATTEMPT] Trying {method_name} conversion...")
                success = await method_func(pdf_path, output_path)

                if success and os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    print(f"[SUCCESS] Conversion completed using {method_name}")
                    return True, method_name
                else:
                    print(f"[FAILED] {method_name} did not produce valid output")

            except Exception as e:
                print(f"[ERROR] {method_name} failed with exception: {e}")
                continue

        # All methods failed
        print("[FAILURE] All conversion methods failed")
        return False, "all_methods_failed"

    else:
        # Use specific method
        method_map = {
            "pdf2docx": (pdf2docx_convert, PDF2DOCX_AVAILABLE, "pdf2docx library"),
            "pymupdf": (pymupdf_to_word, True, "PyMuPDF (fitz)"),
            "pdfplumber": (pdfplumber_to_word_advanced, True, "pdfplumber"),
        }

        if method not in method_map:
            print(f"[ERROR] Invalid method specified: {method}")
            return False, "invalid_method"

        method_func, available, method_description = method_map[method]

        if not available:
            print(f"[ERROR] {method_description} is not available")
            return False, f"{method}_not_available"

        try:
            print(f"[INFO] Using specific method: {method_description}")
            success = await method_func(pdf_path, output_path)

            if success and os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"[SUCCESS] Conversion completed using {method}")
                return True, method
            else:
                print(f"[FAILED] {method} did not produce valid output")
                return False, f"{method}_failed"

        except Exception as e:
            print(f"[ERROR] {method} conversion failed: {e}")
            return False, f"{method}_exception"


# Quality evaluation function for future use
def evaluate_conversion_quality(docx_path: str) -> float:
    """
    Evaluate the quality of a converted Word document

    This function scores a converted document based on:
    - Structure (paragraphs, headings)
    - Formatting (bold, italic, colors)
    - Content (file size, tables)

    Args:
        docx_path: Path to the Word document to evaluate

    Returns:
        float: Quality score (higher is better)

    Note:
        This is a heuristic evaluation and may not reflect actual quality
        in all cases. It's useful for comparing multiple conversion methods.
    """
    score = 0.0

    try:
        from docx import Document

        doc = Document(docx_path)

        # Factor 1: Number of paragraphs (more structure = better)
        paragraph_count = len(doc.paragraphs)
        score += paragraph_count * 0.1

        # Factor 2: Presence of formatting
        formatted_runs = 0
        colored_runs = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold or run.italic or run.underline:
                    formatted_runs += 1
                if run.font.color and run.font.color.rgb:
                    colored_runs += 1

        score += formatted_runs * 0.5
        score += colored_runs * 1.0

        # Factor 3: File size (within reason)
        file_size = os.path.getsize(docx_path)
        # Normalize to reasonable range (10KB - 10MB)
        size_score = min(file_size / 10000, 100)
        score += size_score

        # Factor 4: Tables present
        table_count = len(doc.tables)
        score += table_count * 2.0

        # Factor 5: Images (if any)
        # Note: Counting images in python-docx is complex, skip for now

        print(f"[QUALITY] Document score: {score:.2f}")
        print(f"[QUALITY] - Paragraphs: {paragraph_count}")
        print(f"[QUALITY] - Formatted runs: {formatted_runs}")
        print(f"[QUALITY] - Colored runs: {colored_runs}")
        print(f"[QUALITY] - Tables: {table_count}")
        print(f"[QUALITY] - File size: {file_size} bytes")

    except Exception as e:
        print(f"[WARNING] Quality evaluation failed: {e}")
        return 0.0

    return score


# Compatibility alias for existing code
async def pdf_to_word_converter(pdf_path: str, output_path: str) -> bool:
    """
    Backward compatible function for existing code

    This function maintains compatibility with the existing API
    while using the enhanced converter internally.

    Args:
        pdf_path: Input PDF file path
        output_path: Output DOCX file path

    Returns:
        bool: True if successful, False otherwise
    """
    success, method = await pdf_to_word_converter_enhanced(
        pdf_path, output_path, method="auto"
    )
    return success
