# pymupdf_method.py - PyMuPDF (fitz) conversion with styling preservation
import os

async def pymupdf_to_word(pdf_path: str, output_path: str) -> bool:
    """
    Convert PDF to Word using PyMuPDF with color and styling preservation
    Best for: Color accuracy, font styling, detailed formatting

    Features:
    - Extracts text with complete font information
    - Preserves colors (RGB)
    - Maintains bold/italic formatting
    - Fast processing

    Args:
        pdf_path: Input PDF file path
        output_path: Output DOCX file path

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        import fitz  # PyMuPDF
        from docx import Document
        from docx.shared import RGBColor, Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        print(f"[DEBUG] Starting PyMuPDF conversion: {pdf_path} -> {output_path}")

        # Create Word document
        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Open PDF
        pdf_document = fitz.open(pdf_path)

        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]

            # Add page break for subsequent pages
            if page_num > 0:
                doc.add_page_break()

            # Extract text with detailed formatting using dictionary output
            blocks = page.get_text("dict", flags=11)["blocks"]

            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        paragraph = doc.add_paragraph()

                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if not text.strip():
                                continue

                            font_size = span.get("size", 12)
                            font_name = span.get("font", "Arial")
                            color = span.get("color", 0)  # Integer color
                            flags = span.get("flags", 0)

                            # Convert color from integer to RGB
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF

                            # Add run with formatting
                            run = paragraph.add_run(text)
                            run.font.size = Pt(font_size)
                            run.font.name = font_name

                            # Apply color (only if not black)
                            if color != 0:
                                run.font.color.rgb = RGBColor(r, g, b)

                            # Apply bold/italic based on flags
                            # Flag bit 4: bold, Flag bit 1: italic
                            if flags & (1 << 4):
                                run.bold = True
                            if flags & (1 << 1):
                                run.italic = True

        pdf_document.close()

        # Save Word document
        doc.save(output_path)

        # Verify output
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"[DEBUG] PyMuPDF conversion completed successfully")
            return True
        else:
            print(f"[WARNING] PyMuPDF conversion failed - output file not created")
            return False

    except ImportError as e:
        print(f"[ERROR] Required library not available: {e}")
        return False
    except Exception as e:
        print(f"[WARNING] PyMuPDF conversion failed: {e}")
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass
        return False
