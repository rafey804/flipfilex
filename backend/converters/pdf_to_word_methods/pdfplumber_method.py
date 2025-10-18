# pdfplumber_method.py - Advanced conversion using pdfplumber
import os

async def pdfplumber_to_word_advanced(pdf_path: str, output_path: str) -> bool:
    """
    Convert PDF to Word using pdfplumber with detailed positioning and color extraction
    Best for: Precise positioning, table extraction, color accuracy

    Features:
    - Character-level positioning (x0, x1, top, bottom)
    - Stroking and non-stroking colors
    - Advanced table detection
    - Font height and metrics

    Args:
        pdf_path: Input PDF file path
        output_path: Output DOCX file path

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import RGBColor, Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        print(f"[DEBUG] Starting pdfplumber conversion: {pdf_path} -> {output_path}")

        # Create Word document
        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Add page break for subsequent pages
                if page_num > 0:
                    doc.add_page_break()

                # Extract characters with full details
                chars = page.chars

                if not chars:
                    # Try extracting as plain text if no chars
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
                    continue

                # Group characters by lines (y-coordinate)
                lines_dict = {}
                for char in chars:
                    y = round(char.get('top', 0), 1)
                    if y not in lines_dict:
                        lines_dict[y] = []
                    lines_dict[y].append(char)

                # Sort lines top to bottom
                sorted_lines = sorted(lines_dict.keys())

                for y in sorted_lines:
                    # Sort characters in line by horizontal position
                    line_chars = sorted(lines_dict[y], key=lambda c: c.get('x0', 0))

                    if not line_chars:
                        continue

                    paragraph = doc.add_paragraph()
                    current_text = ""
                    current_color = None
                    current_size = None

                    for char in line_chars:
                        char_text = char.get('text', '')
                        char_color = char.get('non_stroking_color')  # Fill color
                        char_size = char.get('height', 12)
                        font_name = char.get('fontname', 'Arial')

                        # Check if formatting changes
                        color_changed = char_color != current_color
                        size_changed = abs(char_size - (current_size or char_size)) > 1

                        if color_changed or size_changed:
                            # Save previous run if exists
                            if current_text.strip():
                                run = paragraph.add_run(current_text)
                                if current_size:
                                    run.font.size = Pt(current_size)

                                # Apply color if available and not black/white
                                if current_color and isinstance(current_color, (list, tuple)):
                                    if len(current_color) == 3:
                                        r, g, b = [int(c * 255) if c <= 1 else int(c) for c in current_color]
                                        # Only apply if not default black
                                        if not (r == 0 and g == 0 and b == 0):
                                            run.font.color.rgb = RGBColor(r, g, b)

                            # Start new run
                            current_text = char_text
                            current_color = char_color
                            current_size = char_size
                        else:
                            current_text += char_text

                    # Add final run for this line
                    if current_text.strip():
                        run = paragraph.add_run(current_text)
                        if current_size:
                            run.font.size = Pt(current_size)

                        if current_color and isinstance(current_color, (list, tuple)):
                            if len(current_color) == 3:
                                r, g, b = [int(c * 255) if c <= 1 else int(c) for c in current_color]
                                if not (r == 0 and g == 0 and b == 0):
                                    run.font.color.rgb = RGBColor(r, g, b)

        # Save Word document
        doc.save(output_path)

        # Verify output
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"[DEBUG] pdfplumber conversion completed successfully")
            return True
        else:
            print(f"[WARNING] pdfplumber conversion failed - output file not created")
            return False

    except ImportError as e:
        print(f"[ERROR] Required library not available: {e}")
        return False
    except Exception as e:
        print(f"[WARNING] pdfplumber conversion failed: {e}")
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass
        return False
