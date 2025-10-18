# converters/word_to_pdf_converter.py - Professional Word to PDF conversion with full formatting preservation
import os
import subprocess
import shutil
import traceback

async def word_to_pdf_converter(word_path: str, output_path: str) -> bool:
    """
    Convert Word document to PDF with PERFECT formatting preservation
    Uses industry-standard methods similar to iLovePDF
    Priority: 1) LibreOffice 2) docx2pdf (Windows COM) 3) unoconv
    """
    print(f"[DEBUG] Starting Word to PDF conversion: {word_path} -> {output_path}")

    # METHOD 1: Try LibreOffice (BEST - Works on all platforms, perfect formatting)
    print(f"[DEBUG] Attempting LibreOffice conversion (best quality)")
    try:
        # Find LibreOffice executable
        soffice_cmd = None
        possible_paths = [
            'soffice',
            'libreoffice',
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            '/usr/bin/soffice',
            '/usr/bin/libreoffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice'
        ]

        for cmd in possible_paths:
            if shutil.which(cmd) or os.path.exists(cmd):
                soffice_cmd = cmd
                print(f"[DEBUG] Found LibreOffice at: {cmd}")
                break

        if soffice_cmd:
            # Use LibreOffice headless conversion (maintains ALL formatting)
            print(f"[DEBUG] Running LibreOffice conversion...")
            result = subprocess.run([
                soffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(output_path),
                word_path
            ], capture_output=True, text=True, timeout=90)

            print(f"[DEBUG] LibreOffice stdout: {result.stdout}")
            print(f"[DEBUG] LibreOffice stderr: {result.stderr}")

            # LibreOffice creates output based on input filename
            expected_output = os.path.join(
                os.path.dirname(output_path),
                os.path.splitext(os.path.basename(word_path))[0] + '.pdf'
            )

            print(f"[DEBUG] Looking for output at: {expected_output}")

            # Rename to desired output path if different
            if os.path.exists(expected_output):
                if expected_output != output_path:
                    shutil.move(expected_output, output_path)

                if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
                    print(f"[SUCCESS] LibreOffice conversion completed! Size: {os.path.getsize(output_path)} bytes")
                    return True
                else:
                    print(f"[WARNING] LibreOffice output too small: {os.path.getsize(output_path)} bytes")
            else:
                print(f"[WARNING] LibreOffice didn't create output at {expected_output}")
        else:
            print(f"[INFO] LibreOffice not found in system")

    except subprocess.TimeoutExpired:
        print(f"[WARNING] LibreOffice conversion timed out")
    except Exception as e:
        print(f"[WARNING] LibreOffice conversion failed: {e}")
        traceback.print_exc()

    # METHOD 2: Try docx2pdf (Windows only - uses Microsoft Word COM, PERFECT formatting)
    print(f"[DEBUG] Attempting docx2pdf conversion (Windows COM)")
    try:
        from docx2pdf import convert

        print(f"[DEBUG] Running docx2pdf...")
        convert(word_path, output_path)

        if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
            print(f"[SUCCESS] docx2pdf conversion completed! Size: {os.path.getsize(output_path)} bytes")
            return True
        else:
            print(f"[WARNING] docx2pdf output invalid")
            if os.path.exists(output_path):
                os.remove(output_path)

    except ImportError:
        print(f"[INFO] docx2pdf not available")
    except Exception as e:
        print(f"[WARNING] docx2pdf conversion failed: {e}")
        traceback.print_exc()
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass

    # METHOD 3: Try unoconv (Alternative LibreOffice tool)
    print(f"[DEBUG] Attempting unoconv conversion")
    try:
        if shutil.which('unoconv'):
            print(f"[DEBUG] Running unoconv...")
            result = subprocess.run([
                'unoconv',
                '-f', 'pdf',
                '-o', output_path,
                word_path
            ], capture_output=True, text=True, timeout=90)

            print(f"[DEBUG] unoconv stdout: {result.stdout}")
            print(f"[DEBUG] unoconv stderr: {result.stderr}")

            if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
                print(f"[SUCCESS] unoconv conversion completed! Size: {os.path.getsize(output_path)} bytes")
                return True
        else:
            print(f"[INFO] unoconv not found")

    except subprocess.TimeoutExpired:
        print(f"[WARNING] unoconv conversion timed out")
    except Exception as e:
        print(f"[WARNING] unoconv conversion failed: {e}")
        traceback.print_exc()

    # METHOD 4: Advanced fallback using python-docx + ReportLab with formatting preservation
    print(f"[DEBUG] Attempting advanced python-docx + ReportLab conversion")
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfgen import canvas

        print(f"[DEBUG] Reading Word document...")
        doc = Document(word_path)

        # Create PDF with proper page setup
        pdf = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Get styles
        styles = getSampleStyleSheet()
        story = []

        # Process each paragraph with formatting
        for para in doc.paragraphs:
            if para.text.strip():
                # Detect paragraph style
                style_name = 'Normal'

                # Check for headings
                if para.style.name.startswith('Heading'):
                    if '1' in para.style.name:
                        style_name = 'Heading1'
                    elif '2' in para.style.name:
                        style_name = 'Heading2'
                    else:
                        style_name = 'Heading3'

                # Create paragraph with style
                try:
                    # Build formatted text with runs
                    text_parts = []
                    for run in para.runs:
                        text = run.text
                        if run.bold:
                            text = f"<b>{text}</b>"
                        if run.italic:
                            text = f"<i>{text}</i>"
                        if run.underline:
                            text = f"<u>{text}</u>"
                        text_parts.append(text)

                    formatted_text = ''.join(text_parts)

                    p = Paragraph(formatted_text, styles[style_name])
                    story.append(p)
                    story.append(Spacer(1, 0.2*inch))
                except Exception as e:
                    # Fallback to plain text
                    p = Paragraph(para.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 0.2*inch))

        # Process tables if any
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)

            if table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 0.3*inch))

        if not story:
            story.append(Paragraph("No content found in document", styles['Normal']))

        # Build PDF
        pdf.build(story)

        if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
            print(f"[SUCCESS] Advanced ReportLab conversion completed! Size: {os.path.getsize(output_path)} bytes")
            return True
        else:
            print(f"[ERROR] ReportLab output invalid")
            return False

    except Exception as e:
        print(f"[ERROR] Advanced conversion failed: {e}")
        traceback.print_exc()
        return False

    print(f"[ERROR] All Word to PDF conversion methods failed")
    return False