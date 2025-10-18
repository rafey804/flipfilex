# converters/pdf_to_word_converter.py - Professional PDF to Word conversion with full formatting preservation
import os
import subprocess
import shutil
import traceback

async def pdf_to_word_converter(pdf_path: str, output_path: str) -> bool:
    """
    Convert PDF to Word document with PERFECT formatting preservation
    Uses industry-standard methods similar to iLovePDF
    Priority: 1) pdf2docx 2) LibreOffice 3) PyMuPDF with advanced extraction
    """
    print(f"[DEBUG] Starting PDF to Word conversion: {pdf_path} -> {output_path}")

    # METHOD 1: Try pdf2docx (BEST for formatting preservation - specifically designed for PDF to DOCX)
    print(f"[DEBUG] Attempting pdf2docx conversion (best quality for PDF to Word)")
    try:
        from pdf2docx import Converter

        print(f"[DEBUG] Running pdf2docx with advanced settings...")
        cv = Converter(pdf_path)

        # Convert with all formatting options enabled
        cv.convert(output_path, start=0, end=None)
        cv.close()

        if os.path.exists(output_path) and os.path.getsize(output_path) > 5000:
            print(f"[SUCCESS] pdf2docx conversion completed! Size: {os.path.getsize(output_path)} bytes")
            return True
        else:
            print(f"[WARNING] pdf2docx output too small or invalid")
            if os.path.exists(output_path):
                os.remove(output_path)

    except ImportError:
        print(f"[INFO] pdf2docx not available")
    except Exception as e:
        print(f"[WARNING] pdf2docx conversion failed: {e}")
        traceback.print_exc()
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass

    # METHOD 2: Try LibreOffice conversion (Good quality, universal tool)
    print(f"[DEBUG] Attempting LibreOffice conversion")
    try:
        # Find LibreOffice executable
        soffice_cmd = None
        possible_paths = [
            'soffice',
            'libreoffice',
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            '/usr/bin/soffice',
            '/usr/bin/libreoffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice'
        ]

        for cmd in possible_paths:
            if shutil.which(cmd) or os.path.exists(cmd):
                soffice_cmd = cmd
                print(f"[DEBUG] Found LibreOffice at: {cmd}")
                break

        if soffice_cmd:
            print(f"[DEBUG] Running LibreOffice conversion...")
            result = subprocess.run([
                soffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(output_path),
                pdf_path
            ], capture_output=True, text=True, timeout=90)

            print(f"[DEBUG] LibreOffice stdout: {result.stdout}")
            print(f"[DEBUG] LibreOffice stderr: {result.stderr}")

            # LibreOffice creates output based on input filename
            expected_output = os.path.join(
                os.path.dirname(output_path),
                os.path.splitext(os.path.basename(pdf_path))[0] + '.docx'
            )

            print(f"[DEBUG] Looking for output at: {expected_output}")

            # Rename to desired output path if different
            if os.path.exists(expected_output):
                if expected_output != output_path:
                    shutil.move(expected_output, output_path)

                if os.path.exists(output_path) and os.path.getsize(output_path) > 5000:
                    print(f"[SUCCESS] LibreOffice conversion completed! Size: {os.path.getsize(output_path)} bytes")
                    return True
                else:
                    print(f"[WARNING] LibreOffice output too small: {os.path.getsize(output_path)} bytes")
            else:
                print(f"[WARNING] LibreOffice didn't create output at {expected_output}")
        else:
            print(f"[INFO] LibreOffice not found in system")

    except subprocess.TimeoutExpired:
        print(f"[WARNING] LibreOffice conversion timed out")
    except Exception as e:
        print(f"[WARNING] LibreOffice conversion failed: {e}")
        traceback.print_exc()

    # METHOD 3: Try Advanced PyMuPDF (fitz) with detailed formatting extraction
    print(f"[DEBUG] Attempting PyMuPDF conversion with formatting preservation")
    try:
        import fitz  # PyMuPDF
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document()

        # Set narrow margins to match typical PDF layout
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Open PDF with PyMuPDF
        pdf_doc = fitz.open(pdf_path)
        total_text_added = 0

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]

            # Add page break for subsequent pages
            if page_num > 0:
                doc.add_page_break()

            # Extract text blocks with formatting information
            blocks = page.get_text("dict")["blocks"]
            print(f"[DEBUG] Page {page_num}: Found {len(blocks)} blocks")

            for block in blocks:
                if block.get("type") == 0:  # Text block
                    lines = block.get("lines", [])

                    for line in lines:
                        spans = line.get("spans", [])
                        if not spans:
                            continue

                        para = doc.add_paragraph()
                        line_has_content = False

                        for span in spans:
                            text = span.get("text", "")

                            if text:
                                line_has_content = True

                                # Extract formatting information
                                font_size = span.get("size", 11)
                                font_name = span.get("font", "Arial")
                                color = span.get("color", 0)
                                flags = span.get("flags", 0)

                                # Add run with text
                                run = para.add_run(text)
                                total_text_added += len(text)

                                # Apply font size
                                try:
                                    run.font.size = Pt(font_size)
                                except:
                                    pass

                                # Apply bold/italic
                                if "Bold" in font_name or (flags & 2**4):
                                    run.bold = True
                                if "Italic" in font_name or (flags & 2**1):
                                    run.italic = True

                                # Apply color
                                try:
                                    if color != 0:
                                        r = (color >> 16) & 0xFF
                                        g = (color >> 8) & 0xFF
                                        b = color & 0xFF
                                        run.font.color.rgb = RGBColor(r, g, b)
                                except:
                                    pass

                        # Remove empty paragraphs
                        if not line_has_content and para._element.getparent() is not None:
                            para._element.getparent().remove(para._element)

        pdf_doc.close()
        print(f"[DEBUG] Total characters added: {total_text_added}")

        # Save the document
        doc.save(output_path)

        if os.path.exists(output_path) and os.path.getsize(output_path) > 5000 and total_text_added > 0:
            print(f"[SUCCESS] PyMuPDF conversion completed with {total_text_added} characters")
            return True
        else:
            print(f"[WARNING] PyMuPDF created invalid document (chars: {total_text_added})")
            if os.path.exists(output_path):
                os.remove(output_path)

    except Exception as e:
        print(f"[WARNING] PyMuPDF conversion failed: {e}")
        traceback.print_exc()

    # METHOD 4: Try pdfplumber with enhanced formatting
    print(f"[DEBUG] Attempting pdfplumber conversion")
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        total_content = 0
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                if page_num > 0:
                    doc.add_page_break()

                # Extract characters with positioning
                chars = page.chars
                if not chars:
                    continue

                # Group by lines
                lines_dict = {}
                for char in chars:
                    y = round(char['top'], 1)
                    if y not in lines_dict:
                        lines_dict[y] = []
                    lines_dict[y].append(char)

                # Process each line
                for y in sorted(lines_dict.keys()):
                    line_chars = sorted(lines_dict[y], key=lambda c: c['x0'])
                    line_text = ''.join([c['text'] for c in line_chars]).strip()

                    if line_text:
                        para = doc.add_paragraph()

                        # Extract font info
                        font_sizes = [c.get('size', 11) for c in line_chars]
                        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 11

                        run = para.add_run(line_text)
                        run.font.size = Pt(avg_font_size)
                        total_content += len(line_text)

                        # Try to extract color if available
                        if 'stroking_color' in line_chars[0]:
                            color = line_chars[0]['stroking_color']
                            if color and isinstance(color, (list, tuple)) and len(color) >= 3:
                                r, g, b = int(color[0] * 255), int(color[1] * 255), int(color[2] * 255)
                                run.font.color.rgb = RGBColor(r, g, b)

        doc.save(output_path)

        if os.path.exists(output_path) and os.path.getsize(output_path) > 5000 and total_content > 0:
            print(f"[SUCCESS] pdfplumber conversion completed")
            return True
        else:
            print(f"[WARNING] pdfplumber output invalid")
            if os.path.exists(output_path):
                os.remove(output_path)

    except Exception as e:
        print(f"[WARNING] pdfplumber conversion failed: {e}")
        traceback.print_exc()

    # Final fallback to basic PyPDF2
    print(f"[DEBUG] Final fallback to PyPDF2")
    try:
        import PyPDF2
        from docx import Document
        from docx.shared import Inches

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num, page in enumerate(pdf_reader.pages):
                if page_num > 0:
                    doc.add_page_break()

                text = page.extract_text()
                if text.strip():
                    for line in text.split('\n'):
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)

        doc.save(output_path)

        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"[SUCCESS] PyPDF2 conversion completed")
            return True
        else:
            return False

    except Exception as e:
        print(f"[ERROR] All conversion methods failed: {e}")
        traceback.print_exc()
        return False
